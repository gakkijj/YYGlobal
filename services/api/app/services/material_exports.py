import re
from io import BytesIO
from typing import Optional
from urllib.parse import quote

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.models.entities import MaterialDraft, Program

PAGE_WIDTH = 595
PAGE_HEIGHT = 842
MARGIN_X = 58
TOP_Y = 62
BOTTOM_Y = 58
BODY_SIZE = 10.5
BODY_LEADING = 16
FONT_NAME = "china-s"


def _safe_part(value: str, fallback: str) -> str:
    value = re.sub(r"[^A-Za-z0-9]+", "-", value).strip("-")
    return value[:60] or fallback


def export_filename(draft: MaterialDraft, program: Optional[Program], extension: str) -> str:
    material = draft.kind.upper()
    scope = _safe_part(program.university, "Program") if program else "General"
    return f"{material}-{scope}-v{draft.version_number}.{extension}"


def content_disposition(filename: str) -> str:
    ascii_name = _safe_part(filename.rsplit(".", 1)[0], "material")
    extension = filename.rsplit(".", 1)[-1]
    fallback = f"{ascii_name}.{extension}"
    return f"attachment; filename=\"{fallback}\"; filename*=UTF-8''{quote(filename)}"


def text_export(draft: MaterialDraft) -> bytes:
    return draft.content.encode("utf-8")


def _set_run_font(run: object, size: float, bold: bool = False, color: str = "202724") -> None:
    run.font.name = "Arial Unicode MS"
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    fonts.set(qn("w:ascii"), "Arial Unicode MS")
    fonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _add_markdown_runs(paragraph: object, value: str, size: float = 10.5) -> None:
    cursor = 0
    pattern = r"\*\*(.+?)\*\*|__(.+?)__|(?<!\*)\*([^*]+?)\*(?!\*)"
    for match in re.finditer(pattern, value):
        if match.start() > cursor:
            _set_run_font(paragraph.add_run(value[cursor : match.start()]), size)
        text = next(group for group in match.groups() if group is not None)
        run = paragraph.add_run(text)
        _set_run_font(run, size, bold=match.group(3) is None)
        run.italic = match.group(3) is not None
        cursor = match.end()
    if cursor < len(value):
        _set_run_font(paragraph.add_run(value[cursor:]), size)


def _add_page_number(paragraph: object) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    _set_run_font(run, 8, color="66706B")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def docx_export(draft: MaterialDraft, program: Optional[Program]) -> bytes:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72 if draft.kind == "cv" else 0.85)
    section.bottom_margin = Inches(0.72 if draft.kind == "cv" else 0.85)
    section.left_margin = Inches(0.8 if draft.kind == "cv" else 1.0)
    section.right_margin = Inches(0.8 if draft.kind == "cv" else 1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(10.5 if draft.kind == "cv" else 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5 if draft.kind == "cv" else 8)
    normal.paragraph_format.line_spacing = 1.08 if draft.kind == "cv" else 1.2

    for style_name, size, before, after in (
        ("Heading 1", 16, 11, 5),
        ("Heading 2", 13, 9, 4),
        ("Heading 3", 11.5, 7, 3),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(25, 65, 50)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    scope_label = program.university if program else "General"
    header = section.header.paragraphs[0]
    _set_run_font(
        header.add_run(f"{draft.kind.upper()}  ·  {scope_label}  ·  v{draft.version_number}"),
        8,
        color="66706B",
    )
    _add_page_number(section.footer.paragraphs[0])

    for raw_line in draft.content.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw_line.rstrip()
        if line.startswith("### "):
            _add_markdown_runs(document.add_paragraph(style="Heading 3"), line[4:], 11.5)
        elif line.startswith("## "):
            _add_markdown_runs(document.add_paragraph(style="Heading 2"), line[3:], 13)
        elif line.startswith("# "):
            _add_markdown_runs(document.add_paragraph(style="Heading 1"), line[2:], 16)
        elif line.strip() in {"---", "***", "___"}:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(5)
            borders = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            for key, value in (("val", "single"), ("sz", "4"), ("space", "1"), ("color", "D5DDD8")):
                bottom.set(qn(f"w:{key}"), value)
            borders.append(bottom)
            paragraph._p.get_or_add_pPr().append(borders)
        elif line.startswith(("- ", "* ", "• ")):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.35)
            paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            paragraph.paragraph_format.space_after = Pt(3)
            _add_markdown_runs(paragraph, line[2:], 10.5)
        elif line:
            paragraph = document.add_paragraph()
            _add_markdown_runs(paragraph, line, 10.5 if draft.kind == "cv" else 11)
        else:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1)

    document.core_properties.title = draft.title
    document.core_properties.subject = f"YYGlobal {draft.kind.upper()} v{draft.version_number}"
    document.core_properties.author = "YYGlobal"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _pdf_text(value: str) -> str:
    value = value.replace("\u2011", "-").replace("\u2013", "-").replace("\u2014", "-")
    value = re.sub(r"\*\*(.+?)\*\*", r"\1", value)
    value = re.sub(r"__(.+?)__", r"\1", value)
    value = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"\1", value)
    value = re.sub(r"\[([^]]+)]\(([^)]+)\)", r"\1 (\2)", value)
    return value


def _wrap_line(text: str, font: fitz.Font, size: float, width: float) -> list[str]:
    if not text:
        return [""]
    lines: list[str] = []
    current = ""
    for character in text:
        candidate = current + character
        if current and font.text_length(candidate, fontsize=size) > width:
            break_at = current.rfind(" ")
            if break_at > max(2, len(current) // 2):
                lines.append(current[:break_at].rstrip())
                current = current[break_at + 1 :] + character
            else:
                lines.append(current.rstrip())
                current = character
        else:
            current = candidate
    lines.append(current.rstrip())
    return lines


def pdf_export(draft: MaterialDraft, program: Optional[Program]) -> bytes:
    document = fitz.open()
    font = fitz.Font(fontname=FONT_NAME)
    page = document.new_page(width=PAGE_WIDTH, height=PAGE_HEIGHT)
    page.insert_font(fontname="MaterialFont", fontbuffer=font.buffer)
    y = TOP_Y

    def new_page() -> None:
        nonlocal page, y
        page = document.new_page(width=PAGE_WIDTH, height=PAGE_HEIGHT)
        page.insert_font(fontname="MaterialFont", fontbuffer=font.buffer)
        y = TOP_Y

    def add_line(text: str, size: float, leading: float, color: tuple[float, float, float]) -> None:
        nonlocal y
        wrapped = _wrap_line(text, font, size, PAGE_WIDTH - 2 * MARGIN_X)
        for line in wrapped:
            if y + leading > PAGE_HEIGHT - BOTTOM_Y:
                new_page()
            if line:
                page.insert_text(
                    (MARGIN_X, y),
                    line,
                    fontname="MaterialFont",
                    fontsize=size,
                    color=color,
                )
            y += leading

    for raw_line in draft.content.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = _pdf_text(raw_line.rstrip())
        if line.startswith("### "):
            if y + 50 > PAGE_HEIGHT - BOTTOM_Y:
                new_page()
            y += 4
            add_line(line[4:], 12, 18, (0.12, 0.22, 0.18))
        elif line.startswith("## "):
            if y + 55 > PAGE_HEIGHT - BOTTOM_Y:
                new_page()
            y += 7
            add_line(line[3:], 14, 21, (0.08, 0.18, 0.14))
        elif line.startswith("# "):
            if y + 60 > PAGE_HEIGHT - BOTTOM_Y:
                new_page()
            y += 4
            add_line(line[2:], 17, 25, (0.05, 0.12, 0.1))
        elif line.strip() in {"---", "***", "___"}:
            if y + 16 > PAGE_HEIGHT - BOTTOM_Y:
                new_page()
            y += 7
            page.draw_line(
                (MARGIN_X, y),
                (PAGE_WIDTH - MARGIN_X, y),
                color=(0.82, 0.84, 0.83),
                width=0.5,
            )
            y += 9
        elif line.startswith(("- ", "* ")):
            add_line(f"- {line[2:]}", BODY_SIZE, BODY_LEADING, (0.12, 0.12, 0.12))
        elif not line:
            y += BODY_LEADING * 0.55
            if y > PAGE_HEIGHT - BOTTOM_Y:
                new_page()
        else:
            add_line(line, BODY_SIZE, BODY_LEADING, (0.12, 0.12, 0.12))

    material_label = draft.kind.upper()
    scope_label = program.university if program else "General"
    total_pages = len(document)
    for index, current_page in enumerate(document):
        current_page.insert_font(fontname="MaterialFont", fontbuffer=font.buffer)
        current_page.draw_line(
            (MARGIN_X, 35), (PAGE_WIDTH - MARGIN_X, 35), color=(0.75, 0.8, 0.77), width=0.6
        )
        current_page.insert_text(
            (MARGIN_X, 27),
            f"{material_label} · {scope_label} · v{draft.version_number}",
            fontname="MaterialFont",
            fontsize=8,
            color=(0.35, 0.42, 0.39),
        )
        footer = f"{index + 1} / {total_pages}"
        footer_width = font.text_length(footer, fontsize=8)
        current_page.insert_text(
            (PAGE_WIDTH - MARGIN_X - footer_width, PAGE_HEIGHT - 28),
            footer,
            fontname="MaterialFont",
            fontsize=8,
            color=(0.4, 0.4, 0.4),
        )

    document.set_metadata(
        {
            "title": draft.title,
            "subject": f"YYGlobal {material_label} v{draft.version_number}",
            "creator": "YYGlobal",
        }
    )
    output = BytesIO()
    document.save(output, garbage=4, deflate=True)
    document.close()
    return output.getvalue()
